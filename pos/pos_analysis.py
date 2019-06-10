# -*- coding: utf-8 -*-
"""
Created on Wed Apr 10 10:44:10 2019

@author: Guilherme Neumann
"""
import os
import stats
import seaborn as sns
import pandas as pd
from scipy import stats as stat
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt 
import random

pd.set_option('display.max_columns', None)
pd.set_option('display.width', 200)

def replace(group, stds):
    group[np.abs(group - group.mean()) > stds * group.std()] = np.nan
    return group

def table_docx(document,df):
    t = document.add_table(df.shape[0]+1, df.shape[1]+1,style='Light Grid Accent 1')
    for j in range(df.shape[-1]):
        t.cell(0,j+1).text = df.columns[j]
    for i in range (len(df.index)):
        t.cell(i+1,0).text=df.index[i]
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j+1).text = str(df.values[i,j])
def tests(var,data,document,met,colors,run=''):
    df=data.copy()
    pk_reads=1
    pk_assemblers=1
    if var!="Ecoli strains":
        if var!='Phred':
            a = df.index.to_series().str.rsplit('_').str[-1].astype(float).sort_values()
            df = df.reindex(index=a.index)
    else:
        d={"Genome":[],"Assembler":[],met:[]}
        df_print=pd.DataFrame(data=d)
        for col in df.T.columns:
            gen=col.split('-')
            gen=gen[0][4:]
            for i,row in enumerate(df.T[col]):
                df_print=df_print.append(pd.DataFrame([[gen,df.T.index[i],row]],columns=["Genome","Assembler",met]), ignore_index=True)
        sns.swarmplot(x="Genome",y=met,hue="Assembler",data=df_print,palette="Set1")
        plt.ylabel(met)
        plt.xticks(rotation=90)
        #5132068,4641652,5437407,4894879
        #iai39,k12,o104,o83
        fig = plt.gcf()
        fig.set_size_inches(12, 8)
        fig.savefig(out+met+'_swarm'+run+'.png', dpi=300)
        fig.clear()
        document.add_picture(out+met+'_swarm'+run+'.png', width=Inches(6.25))
    table_docx(document,df)

    ob=stats.Statistics(df,exp,var,met,"outputs_csv/",samples_int,run)
    if var!="Ecoli strains" and run!='without_zscore':
        if not os.path.isfile(out+met+run+".png"):
            ob.scatter_plot(colors)
        document.add_picture(out+met+run+'.png', width=Inches(6.25))
        
    norms=ob.normality()
    equal=[]
    for item in df.values:
        for each2 in item:
            if each2 not in equal:
                equal.append(each2)
        #check if all the results are not the same. At least one need to be different    
    if len(equal)>1:
        if norms[0] and norms[1]:
            anova=ob.oneway()
            '''anova_assemblers=anova[0].pvalue
            anova_reads=anova[1].pvalue
            if anova_assemblers <0.05 or  anova_reads<0.05:
                #ob.tukey_test()'''
            k=ob.kruskal()
            pk_assemblers=k[0].pvalue
            pk_reads=k[1].pvalue
            if pk_assemblers <0.05 or  pk_reads<0.05:
                nem=ob.nemenyi_test()

        else:
            k=ob.kruskal()
            pk_assemblers=k[0].pvalue
            pk_reads=k[1].pvalue
            if pk_assemblers <0.05 or  pk_reads<0.05:
                nem=ob.nemenyi_test()
                if run!='without_zscore' and var!="Ecoli strains":
                    ob.correlation_pearson()
                    document.add_picture(out+'corr_'+met+run+'.png', width=Inches(6.00))
                    try:
                        ob.linear_regression()                        
                    except:
                        pass                    
    ob.results.close()    
            
    temp=open("outputs_csv/final_stats_"+met+""+run+".txt")
    for line in temp:
        document.add_paragraph(line.strip())
    if pk_assemblers <0.05: 
        document.add_heading("Nemenyi Assemblers",level=5)
        t = document.add_table(nem[0].shape[0]+1, nem[0].shape[1]+1,style='Light Grid Accent 1')
        for j in range(nem[0].shape[-1]):
            t.cell(0,j+1).text = nem[0].columns[j]
        for i in range (len(nem[0].index)):
            t.cell(i+1,0).text=nem[0].index[i]
        # add the rest of the data frame
        for i in range(nem[0].shape[0]):
            for j in range(nem[0].shape[-1]):
                t.cell(i+1,j+1).text = str(nem[0].values[i,j])
    if pk_reads<0.05: 
        document.add_heading("Nemenyi samples",level=5)
        t = document.add_table(nem[1].shape[0]+1, nem[1].shape[1]+1,style='Light Grid Accent 1')
        for j in range(nem[1].shape[-1]):
            t.cell(0,j+1).text = nem[1].columns[j]
        for i in range (len(nem[1].index)):
            t.cell(i+1,0).text=nem[1].index[i]
        # add the rest of the data frame
        for i in range(nem[1].shape[0]):
            for j in range(nem[1].shape[-1]):
                t.cell(i+1,j+1).text = str(nem[1].values[i,j])
    if run!='without_zscore' and var!="Ecoli strains":
        df.index=samples_int
        if not os.path.isfile(out+met+'_plot'+run+'.png'):
            df[var]=samples
            df.plot(x=var, title=met)
            plt.ylabel(met)
            plt.xticks(rotation=90)
                
            fig = plt.gcf()
            fig.set_size_inches(12, 8)
            fig.savefig(out+met+'_plot'+run+'.png', dpi=300)
            fig.clear()
        document.add_picture(out+met+'_plot'+run+'.png', width=Inches(6.25))
            
    if not os.path.isfile(out+met+'_box'+run+'.png'):
        df.plot.box(title=met,whis=3)
        plt.ylabel(met)
        plt.xticks(rotation=90)
        
        fig = plt.gcf()
        fig.set_size_inches(12, 8)
        fig.savefig(out+met+'_box'+run+'.png', dpi=300)
        fig.clear()
            
    document.add_picture(out+met+'_box'+run+'.png', width=Inches(6.25))
    return ob

exp=input("Experiment Name: ")
var=input("Type the variable of the Experiment: ")
samples=input("and the samples variation separated by comma (e.g. 100,150,200,250,300):")
document = Document()
document.add_heading("Complete Report "+exp, 0)

samples=samples.split(",")
samples_int=[]
if var!='Ecoli strains':
    for s in samples:
        samples_int.append(float(s))
else:
    samples_int=samples
metrics={}
outs=os.listdir("outputs_csv")
out="outputs_csv/"
colors=[]
for each in outs:    
    
    met=each.split(".")
    if met[1]=="csv" and "corr_" not in met[0] and "nem_" not in met[0]:
        df=pd.read_csv(out+each, header=0, index_col=0)
        document.add_heading(met[0], level=1)
        #df.columns=['0.00010','0.00032','0.001','0.0031','0.01','0.031','0.1','0.31','0.79']
        df.columns=['Q40','Q35','Q30','Q25','Q20','Q15','Q10','Q5','Q1']
        #40,35,30,25,20,15,10,5,1
        if var=='Ecoli strains' and met[0] in ("Largest contig","N50","N75"):
            table_docx(document,df)
            df2=pd.read_csv(out+"Total length.csv", header=0, index_col=0)
            for col in df2.columns:
                for i,row in enumerate(df2[col]):
                    df[col][i]=(df[col][i]*100)/row
            document.add_heading(met[0]+" normalized - frequency of total length", level=5)
            met[0]=met[0]+" %"
        elif var=='Ecoli strains' and met[0] in ("Largest alignment","NA50","NA75"):
            table_docx(document,df)
            df2=pd.read_csv(out+"Total aligned length.csv", header=0, index_col=0)
            for col in df2.columns:
                for i,row in enumerate(df2[col]):
                    df[col][i]=(df[col][i]*100)/row
            document.add_heading(met[0]+" normalized - frequency of total aligned length", level=5)
            met[0]=met[0]+" %"
        elif var=='Ecoli strains' and met[0] in ("Total aligned length","Total length","NG50"):
            tamanhos=[5132068,4641652,5437407,4894879]
            table_docx(document,df)
            for j,col in enumerate(df.columns):
                for i,row in enumerate(df[col]):
                    df[col][i]=(row*100)/tamanhos[j]
            document.add_heading(met[0]+" normalized - frequency of reference genome length", level=5)
            met[0]=met[0]+" %"
        elif var=='Ecoli strains' and met[0]== "GC (%)":
            gc=[50.6,50.8,50.7,50.7]
            table_docx(document,df)
            for j,col in enumerate(df2.columns):
                for i,row in enumerate(df[col]):
                    df[col][i]=gc[j]-row
            document.add_heading(met[0]+" normalized - difference of reference GC", level=5)
            met[0]="Difference in GC"
        elif var=='Ecoli strains' and met[0]== "# predicted genes (unique)":
            gc=[5092,4566,5081,4532]
            table_docx(document,df)
            for j,col in enumerate(df2.columns):
                for i,row in enumerate(df[col]):
                    df[col][i]=gc[j]-row
            document.add_heading(met[0]+" normalized - difference of reference predicted genes", level=5)
            met[0]="# Genes not predicted"
        metrics[met[0]]=df.T #for each metric in the dir, we create a dataframe
        if colors==[]:
            colors={assembler:["#"+''.join([random.choice('0123456789ABCDEF') for j in range(6)]),random.choice('8*hPov^<>')] for assembler in df.index} #it creates a dict with color and marker type to each assembler, in order to keep the same pattern in all graphs

        ob=tests(var,metrics[met[0]],document,met[0],colors)
        try:
            document.add_heading("Removing Velvet from data:", level=2)
            print("Removing SSAKE from data:")
            #metrics[met[0]]=metrics[met[0]].drop(['ssake'],axis=1)
            metrics[met[0]]=metrics[met[0]].drop(['velvet'],axis=1)
            ob=tests(var,metrics[met[0]],document,met[0],colors,'without_ssake')       
        except:
            pass 
           
        try:
            z=np.abs(stat.zscore(metrics[met[0]]))

            if np.where(z>3)[0]!=[]:
                document.add_heading("Removing outliers through Zscore", level=2)
                print("Removing outliers through Zscore")
                metrics[met[0]]=metrics[met[0]][(z<3).all(axis=1)] 
                document.add_paragraph(str(np.where(z>3)), style='List Bullet')
                ob=tests(var,metrics[met[0]],document,met[0],colors,'without_zscore')                      
        except:
            pass 
        
            
        document.save(exp+".docx")

 
        